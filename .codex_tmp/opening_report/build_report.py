from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


TITLE = "面向成本优化的长期记忆写入阶段模型路由研究"
REFERENCE = Path(r"W:\附件1-3：硕士生开题报告模板0604.docx")
SOURCE_MD = Path(r"S:\Workfile\InfoBudget2\deliverables\硕士学位论文开题报告_初稿.md")
OUTPUT = Path(r"S:\Workfile\InfoBudget2\deliverables\硕士学位论文开题报告_初稿.docx")
ROUTE_PNG = Path(r"S:\Workfile\InfoBudget2\.codex_tmp\opening_report\technical_route.png")


def set_run_font(run, cn="宋体", en="Times New Roman", size=12, bold=None):
    run.font.name = en
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cn)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), en)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), en)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def format_body(p, *, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(24) if indent else Pt(0)
    pf.keep_together = False
    for run in p.runs:
        set_run_font(run, size=size)


def format_heading(p, level):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.keep_with_next = True
    pf.keep_together = True
    pf.line_spacing = 1.0
    if level == 1:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        size = 14
    else:
        pf.space_before = Pt(9)
        pf.space_after = Pt(3)
        size = 12
    for run in p.runs:
        set_run_font(run, cn="黑体", en="Times New Roman", size=size, bold=True)


def add_after(parent, element, before):
    parent.insert(parent.index(before), element)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def new_decimal_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "space")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "480")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def draw_route(path):
    w, h = 1700, 720
    img = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(img)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    bold_path = Path(r"C:\Windows\Fonts\msyhbd.ttc")
    font = ImageFont.truetype(str(font_path), 34)
    title_font = ImageFont.truetype(str(bold_path if bold_path.exists() else font_path), 38)
    boxes = [
        ("长对话历史", 65, 120),
        ("主题分段", 465, 120),
        ("片段与模型表征", 865, 120),
        ("质量预测", 1265, 120),
        ("预算约束分配", 1265, 440),
        ("分档事实抽取", 865, 440),
        ("持久记忆库", 465, 440),
        ("检索问答与评测", 65, 440),
    ]
    bw, bh = 300, 115
    for text, x, y in boxes:
        draw.rounded_rectangle((x, y, x + bw, y + bh), radius=24, fill="#EAF2F8", outline="#2E6F95", width=4)
        bbox = draw.textbbox((0, 0), text, font=font)
        draw.text((x + (bw - (bbox[2] - bbox[0])) / 2, y + (bh - (bbox[3] - bbox[1])) / 2 - 4), text, fill="#17324D", font=font)

    def arrow(a, b):
        draw.line((a[0], a[1], b[0], b[1]), fill="#2E6F95", width=6)
        import math
        ang = math.atan2(b[1] - a[1], b[0] - a[0])
        for off in (2.55, -2.55):
            p = (b[0] + 24 * math.cos(ang + off), b[1] + 24 * math.sin(ang + off))
            draw.line((b[0], b[1], p[0], p[1]), fill="#2E6F95", width=6)

    for i in range(3):
        x, y = boxes[i][1], boxes[i][2]
        nx, ny = boxes[i + 1][1], boxes[i + 1][2]
        arrow((x + bw, y + bh / 2), (nx - 16, ny + bh / 2))
    arrow((boxes[3][1] + bw / 2, boxes[3][2] + bh), (boxes[4][1] + bw / 2, boxes[4][2] - 16))
    for i in range(4, 7):
        x, y = boxes[i][1], boxes[i][2]
        nx, ny = boxes[i + 1][1], boxes[i + 1][2]
        arrow((x, y + bh / 2), (nx + bw + 16, ny + bh / 2))
    draw.text((65, 42), "面向成本优化的长期记忆写入模型路由", fill="#17324D", font=title_font)
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, dpi=(180, 180))


def iter_body(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "p", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "t", Table(child, doc)


def fill_cover(doc):
    blocks = list(iter_body(doc))
    for idx, (kind, block) in enumerate(blocks):
        if kind == "p" and block.text.strip() == "题 目：":
            for kind2, candidate in blocks[idx + 1 :]:
                if kind2 == "t":
                    break
                if kind2 == "p":
                    candidate.text = TITLE
                    candidate.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    candidate.paragraph_format.space_before = Pt(4)
                    candidate.paragraph_format.space_after = Pt(0)
                    for run in candidate.runs:
                        set_run_font(run, cn="华文中宋", en="Times New Roman", size=18, bold=True)
                    # The title makes this formerly blank slot slightly taller. Remove
                    # one adjacent blank spacer so the template's explicit page break
                    # remains on the cover rather than spilling to a blank second page.
                    cidx = blocks.index((kind2, candidate))
                    for kind3, spacer in blocks[cidx + 1 :]:
                        if kind3 == "t":
                            break
                        if kind3 == "p" and not spacer.text.strip():
                            spacer._p.getparent().remove(spacer._p)
                            break
                    break
            break
    cover_table = doc.tables[0]
    values = ["[待填写]", "[待填写]", "[待填写]", "[待填写]", "计算机科学与技术学院"]
    for row, value in zip(cover_table.rows, values):
        cell = row.cells[1]
        cell.text = value
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, cn="宋体", en="Times New Roman", size=12, bold=value == "计算机科学与技术学院")


def parse_markdown(path):
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            blocks.append(("h1", line[3:].strip()))
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h2", line[4:].strip()))
            i += 1
            continue
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
            continue
        if re.match(r"^\d+\.\s+", line):
            blocks.append(("number", re.sub(r"^\d+\.\s+", "", line)))
            i += 1
            continue
        if line == "【技术路线图插入位置】":
            blocks.append(("route", ""))
            i += 1
            continue
        blocks.append(("p", line))
        i += 1
    return blocks


def remove_template_body(doc):
    body = doc.element.body
    start = None
    sign_tbl = doc.tables[-1]._tbl
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            if p.text.startswith("课题背景、目的和意义"):
                start = child
                break
    if start is None:
        raise RuntimeError("Template body start marker not found")
    deleting = False
    for child in list(body.iterchildren()):
        if child is start:
            deleting = True
        if child is sign_tbl:
            deleting = False
        if deleting:
            body.remove(child)
    return sign_tbl


def add_schedule_table(doc, rows, sign_tbl):
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=10.5, bold=r_idx == 0)
            if r_idx == 0:
                shade_cell(cell, "D9EAF2")
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1900, 4300, 2106])
    table._tbl.getparent().remove(table._tbl)
    add_after(doc.element.body, table._tbl, sign_tbl)


def add_content(doc, blocks, sign_tbl):
    parent = doc.element.body
    in_refs = False
    number_counter = 0
    previous_kind = None
    first_h1 = True
    for kind, content in blocks:
        if kind != "number":
            number_counter = 0
        if kind == "h1":
            in_refs = content.startswith("7 ")
            p = doc.add_paragraph(style="Title")
            p.add_run(content)
            format_heading(p, 1)
            if first_h1:
                p.paragraph_format.page_break_before = False
                first_h1 = False
            p._p.getparent().remove(p._p)
            add_after(parent, p._p, sign_tbl)
        elif kind == "h2":
            p = doc.add_paragraph(style="Subtitle")
            p.add_run(content)
            format_heading(p, 2)
            p._p.getparent().remove(p._p)
            add_after(parent, p._p, sign_tbl)
        elif kind == "table":
            add_schedule_table(doc, content, sign_tbl)
        elif kind == "route":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            p.add_run().add_picture(str(ROUTE_PNG), width=Inches(5.65))
            p._p.getparent().remove(p._p)
            add_after(parent, p._p, sign_tbl)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(6)
            run = cap.add_run("图1  面向成本优化的长期记忆写入模型路由技术路线")
            set_run_font(run, cn="宋体", en="Times New Roman", size=10.5)
            cap._p.getparent().remove(cap._p)
            add_after(parent, cap._p, sign_tbl)
        elif kind == "number":
            if previous_kind != "number":
                number_counter = 0
            number_counter += 1
            p = doc.add_paragraph()
            p.add_run(f"{number_counter}.  {content}")
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.first_line_indent = Pt(-18)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=12)
            p._p.getparent().remove(p._p)
            add_after(parent, p._p, sign_tbl)
        else:
            p = doc.add_paragraph()
            p.add_run(content)
            if in_refs and re.match(r"^\[\d+\]", content):
                p.paragraph_format.left_indent = Pt(24)
                p.paragraph_format.first_line_indent = Pt(-24)
                p.paragraph_format.line_spacing = 1.25
                p.paragraph_format.space_after = Pt(3)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    set_run_font(run, size=10.5)
            elif content.startswith("q̂("):
                format_body(p, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
                for run in p.runs:
                    set_run_font(run, cn="Cambria Math", en="Cambria Math", size=12)
            else:
                format_body(p)
            p._p.getparent().remove(p._p)
            add_after(parent, p._p, sign_tbl)
        previous_kind = kind
    page = doc.add_paragraph()
    page.add_run().add_break(WD_BREAK.PAGE)
    page._p.getparent().remove(page._p)
    add_after(parent, page._p, sign_tbl)


def style_signature_table(doc):
    table = doc.tables[-1]
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, top=120, bottom=120)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, cn="黑体", en="Times New Roman", size=12, bold=True)


def set_image_alt_text(doc):
    alt_texts = [
        "华中科技大学校徽",
        "面向成本优化的长期记忆写入模型路由技术路线图",
    ]
    for index, doc_pr in enumerate(doc.element.iter(qn("wp:docPr"))):
        alt = alt_texts[index] if index < len(alt_texts) else f"开题报告插图{index + 1}"
        doc_pr.set("title", alt)
        doc_pr.set("descr", alt)


def main():
    if not REFERENCE.exists():
        raise SystemExit(f"Missing template: {REFERENCE}")
    draw_route(ROUTE_PNG)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = Document(OUTPUT)
    fill_cover(doc)
    sign_tbl = remove_template_body(doc)
    blocks = parse_markdown(SOURCE_MD)
    add_content(doc, blocks, sign_tbl)
    style_signature_table(doc)
    set_image_alt_text(doc)
    props = doc.core_properties
    props.title = TITLE
    props.subject = "硕士学位论文开题报告"
    props.keywords = "长期记忆；模型路由；成本优化；预算约束；大语言模型"
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
