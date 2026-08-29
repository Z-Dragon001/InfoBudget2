from __future__ import annotations

import json
import hashlib
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def font_name(run):
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return run.font.name
    return (
        rpr.rFonts.get(qn("w:eastAsia"))
        or rpr.rFonts.get(qn("w:ascii"))
        or rpr.rFonts.get(qn("w:hAnsi"))
        or run.font.name
    )


def para_info(p, idx):
    pf = p.paragraph_format
    runs = []
    for r in p.runs:
        if not r.text:
            continue
        runs.append(
            {
                "text": r.text,
                "font": font_name(r),
                "size_pt": r.font.size.pt if r.font.size else None,
                "bold": r.bold,
                "italic": r.italic,
            }
        )
    return {
        "index": idx,
        "style": p.style.name if p.style else None,
        "text": p.text,
        "alignment": str(p.alignment),
        "space_before_pt": pf.space_before.pt if pf.space_before else None,
        "space_after_pt": pf.space_after.pt if pf.space_after else None,
        "line_spacing": str(pf.line_spacing),
        "first_line_indent_pt": pf.first_line_indent.pt if pf.first_line_indent else None,
        "left_indent_pt": pf.left_indent.pt if pf.left_indent else None,
        "page_break_before": pf.page_break_before,
        "break_types": [b.get(qn("w:type")) or "line" for b in p._p.findall(".//" + qn("w:br"))],
        "runs": runs,
    }


def main():
    src = Path(sys.argv[1])
    doc = Document(src)
    out = {
        "paragraphs": [para_info(p, i) for i, p in enumerate(doc.paragraphs)],
        "tables": [],
        "styles": [],
        "sections": [],
        "headers": [],
        "footers": [],
        "package_parts": [],
        "body_order": [],
    }
    for ti, table in enumerate(doc.tables):
        out["tables"].append(
            {
                "index": ti,
                "style": table.style.name if table.style else None,
                "rows": [
                    [
                        {
                            "text": cell.text,
                            "paragraphs": [para_info(p, pi) for pi, p in enumerate(cell.paragraphs)],
                        }
                        for cell in row.cells
                    ]
                    for row in table.rows
                ],
            }
        )
    para_ids = {id(p._p): i for i, p in enumerate(doc.paragraphs)}
    table_ids = {id(t._tbl): i for i, t in enumerate(doc.tables)}
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = Paragraph(child, doc)
            out["body_order"].append({"kind": "paragraph", "index": para_ids.get(id(child)), "text": p.text})
        elif child.tag == qn("w:tbl"):
            t = Table(child, doc)
            out["body_order"].append({"kind": "table", "index": table_ids.get(id(child)), "text": " | ".join(c.text for c in t.rows[0].cells) if t.rows else ""})
    for s in doc.styles:
        if s.type == 1:
            out["styles"].append(
                {
                    "name": s.name,
                    "font": s.font.name,
                    "size_pt": s.font.size.pt if s.font.size else None,
                    "bold": s.font.bold,
                    "italic": s.font.italic,
                }
            )
    for sec in doc.sections:
        out["sections"].append(
            {
                "width_in": sec.page_width.inches,
                "height_in": sec.page_height.inches,
                "left_in": sec.left_margin.inches,
                "right_in": sec.right_margin.inches,
                "top_in": sec.top_margin.inches,
                "bottom_in": sec.bottom_margin.inches,
                "header_in": sec.header_distance.inches,
                "footer_in": sec.footer_distance.inches,
                "different_first": sec.different_first_page_header_footer,
            }
        )
        out["headers"].append([p.text for p in sec.header.paragraphs])
        out["footers"].append([p.text for p in sec.footer.paragraphs])
    with zipfile.ZipFile(src) as zf:
        for info in sorted(zf.infolist(), key=lambda item: item.filename):
            data = zf.read(info.filename)
            out["package_parts"].append(
                {
                    "path": info.filename,
                    "size": len(data),
                    "sha256": hashlib.sha256(data).hexdigest(),
                }
            )
    payload = json.dumps(out, ensure_ascii=False, indent=2)
    if len(sys.argv) > 2:
        Path(sys.argv[2]).write_text(payload, encoding="utf-8")
    else:
        print(payload)


if __name__ == "__main__":
    main()
